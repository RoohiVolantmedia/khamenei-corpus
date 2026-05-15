"""توابع صادرات: CSV، Excel، JSON، Markdown، DOCX"""
import io
import json
import csv
from datetime import datetime
import pandas as pd


def to_csv_bytes(rows: list[dict]) -> bytes:
    if not rows:
        return b''
    buf = io.StringIO()
    fieldnames = ['doc_id', 'date_persian', 'analytical_tier', 'content_source',
                  'period_label', 'title', 'word_count', 'url',
                  'tag_voice_type', 'tag_form_genre', 'tag_audience',
                  'tag_occasion', 'tag_topics', 'tag_regions', 'tag_tone', 'snippet']
    writer = csv.DictWriter(buf, fieldnames=[f for f in fieldnames if f in rows[0]],
                            extrasaction='ignore')
    writer.writeheader()
    writer.writerows(rows)
    return ('﻿' + buf.getvalue()).encode('utf-8')


def to_excel_bytes(rows: list[dict], keyword: str = '', filters_summary: str = '') -> bytes:
    try:
        from openpyxl import Workbook
        from openpyxl.styles import Font, Alignment, PatternFill
    except ImportError:
        return b''

    wb = Workbook()

    # Sheet 1 — اطلاعات کلی
    ws1 = wb.active
    ws1.title = 'اطلاعات کلی'
    ws1.sheet_view.rightToLeft = True
    ws1.append(['جستجو', keyword or '(بدون کلیدواژه)'])
    ws1.append(['فیلترها', filters_summary])
    ws1.append(['تعداد نتایج', len(rows)])
    ws1.append(['تاریخ صادرات', datetime.now().strftime('%Y-%m-%d %H:%M')])

    # Sheet 2 — متادیتا
    ws2 = wb.create_sheet('متادیتا')
    ws2.sheet_view.rightToLeft = True
    meta_cols = ['doc_id', 'date_persian', 'analytical_tier', 'content_source',
                 'period_label', 'title', 'word_count', 'url',
                 'tag_voice_type', 'tag_form_genre', 'tag_audience',
                 'tag_occasion', 'tag_topics', 'tag_regions', 'tag_tone']
    _write_sheet(ws2, rows, meta_cols)

    # Sheet 3 — متن کامل
    ws3 = wb.create_sheet('متن کامل')
    ws3.sheet_view.rightToLeft = True
    _write_sheet(ws3, rows, ['doc_id', 'date_persian', 'title', 'full_text'],
                 wrap_last=True)

    # Sheet 4 — تگ‌های ۷ بعدی
    ws4 = wb.create_sheet('تگ ۷ بعدی')
    ws4.sheet_view.rightToLeft = True
    tag_cols = ['doc_id', 'tag_voice_type', 'tag_form_genre', 'tag_audience',
                'tag_occasion', 'tag_topics', 'tag_regions', 'tag_tone']
    _write_sheet(ws4, rows, tag_cols)

    # Sheet 5 — کتابخانه‌ی نقل‌قول (خالی برای پر کردن)
    ws5 = wb.create_sheet('نقل‌قول‌ها')
    ws5.sheet_view.rightToLeft = True
    ws5.append(['doc_id', 'تاریخ', 'عنوان', 'نقل‌قول', 'یادداشت', 'تگ'])

    for ws in [ws1, ws2, ws3, ws4, ws5]:
        for cell in ws[1]:
            cell.font = Font(bold=True, name='Vazirmatn')
            cell.fill = PatternFill('solid', fgColor='E8E8E8')

    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _write_sheet(ws, rows, cols, wrap_last=False):
    from openpyxl.styles import Font, Alignment
    available = [c for c in cols if any(c in r for r in rows[:1])]
    ws.append(available)
    for row in rows:
        ws.append([row.get(c, '') for c in available])
    for row in ws.iter_rows(min_row=2):
        for cell in row:
            cell.font = Font(name='Vazirmatn', size=9)
            cell.alignment = Alignment(wrap_text=True, vertical='top')


def to_markdown(doc: dict) -> str:
    lines = [
        f"# {doc.get('title', '(بدون عنوان)')}",
        f"",
        f"**تاریخ:** {doc.get('date_persian', '')}  ",
        f"**منبع:** {doc.get('content_source', '')}  ",
        f"**Tier:** {doc.get('analytical_tier', '')}  ",
        f"**doc_id:** `{doc.get('doc_id', '')}`  ",
        f"**لینک:** {doc.get('url', '')}  ",
        f"",
        f"---",
        f"",
        f"## تگ‌های ۷ بعدی",
        f"- **voice_type:** {doc.get('tag_voice_type','')}",
        f"- **form_genre:** {doc.get('tag_form_genre','')}",
        f"- **audience:** {doc.get('tag_audience','')}",
        f"- **occasion:** {doc.get('tag_occasion','')}",
        f"- **topics:** {doc.get('tag_topics','')}",
        f"- **regions:** {doc.get('tag_regions','')}",
        f"- **tone:** {doc.get('tag_tone','')}",
        f"",
        f"---",
        f"",
        f"## متن",
        f"",
        doc.get('full_text', ''),
        f"",
        f"---",
        f"*صادرشده از پیکره‌ی خامنه‌ای — {datetime.now().strftime('%Y-%m-%d')}*",
    ]
    return '\n'.join(lines)


def to_docx_bytes(doc: dict) -> bytes:
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return b''

    d = Document()
    d.add_heading(doc.get('title', '(بدون عنوان)'), 0)

    meta = d.add_paragraph()
    meta.add_run(f"تاریخ: {doc.get('date_persian','')} | Tier: {doc.get('analytical_tier','')} | doc_id: {doc.get('doc_id','')}")
    meta.add_run(f"\nلینک: {doc.get('url','')}")

    d.add_heading('تگ‌های ۷ بعدی', 2)
    tags = d.add_paragraph()
    for k, v in [('voice_type', 'tag_voice_type'), ('form_genre', 'tag_form_genre'),
                  ('audience', 'tag_audience'), ('occasion', 'tag_occasion'),
                  ('topics', 'tag_topics'), ('regions', 'tag_regions'), ('tone', 'tag_tone')]:
        tags.add_run(f"{k}: ").bold = True
        tags.add_run(str(doc.get(v, '')) + '  \n')

    d.add_heading('متن', 2)
    d.add_paragraph(doc.get('full_text', ''))

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def filename(ext: str) -> str:
    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    return f"khamenei_{ts}.{ext}"
