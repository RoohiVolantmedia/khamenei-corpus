"""صفحه‌ی جستجو"""
import sys, io, csv, json, re as _re, html as _html
from pathlib import Path
from datetime import datetime
sys.path.insert(0, str(Path(__file__).parent.parent))

import streamlit as st
import sqlite3
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from utils.worksets import list_worksets, save_workset, delete_workset, get_workset
from utils.theme import render_header, render_footer

# مسیر دیتابیس — محلی یا سرور
def _resolve_db():
    from pathlib import Path as _P
    local = _P(__file__).resolve().parent.parent.parent / "database.db"
    if local.exists():
        return str(local)
    from utils.db import _get_db_path
    return str(_get_db_path())
DB_PATH = _resolve_db()
FONT_PATH = str(Path(__file__).parent.parent / "assets" / "Vazirmatn.ttf")

render_header("جستجوی پیکره")

# ─── DB ──────────────────────────────────────────────────────────────────────
@st.cache_resource
def get_db():
    c = sqlite3.connect(DB_PATH, check_same_thread=False)
    c.execute("PRAGMA journal_mode=WAL")
    c.row_factory = sqlite3.Row
    return c
conn = get_db()

@st.cache_data(ttl=600)
def load_filter_opts():
    c = sqlite3.connect(DB_PATH); c.row_factory = sqlite3.Row
    def q(sql): return [r[0] for r in c.execute(sql).fetchall()]
    opts = {
        'tiers':   ['tier1','tier2','tier3','tier4'],
        'sources': ['speech','speech_supplement','message','decree'],
        'voice':   q("SELECT DISTINCT tag_voice_type FROM doc_tags WHERE tag_voice_type IS NOT NULL ORDER BY 1"),
        'genre':   q("SELECT DISTINCT tag_form_genre FROM doc_tags WHERE tag_form_genre IS NOT NULL ORDER BY 1"),
        'audience':q("SELECT DISTINCT value FROM doc_tags, json_each(tag_audience) ORDER BY value"),
        'occasion':q("SELECT DISTINCT tag_occasion FROM doc_tags WHERE tag_occasion IS NOT NULL AND tag_occasion!='' ORDER BY 1"),
        'topics':  q("SELECT DISTINCT value FROM doc_tags, json_each(tag_topics) ORDER BY value"),
        'regions': q("SELECT DISTINCT value FROM doc_tags, json_each(tag_regions) ORDER BY value"),
        'tone':    q("SELECT DISTINCT value FROM doc_tags, json_each(tag_tone) ORDER BY value"),
        'periods': [r[0] for r in c.execute("SELECT period_label, MIN(date_persian) f FROM documents WHERE period_label IS NOT NULL GROUP BY period_label ORDER BY f").fetchall()],
    }
    c.close(); return opts

opts = load_filter_opts()

# ─── ترجمه‌های فارسی ──────────────────────────────────────────────────────────
TIER_FA = {
    'tier1': 'tier1 — سخنرانی مستقیم',
    'tier2': 'tier2 — متن ویرایش‌شده',
    'tier3': 'tier3 — خلاصه / گزیده',
    'tier4': 'tier4 — گزارش راوی',
}
VOICE_FA = {
    'primary':   'primary (خامنه‌ای)',
    'secondary': 'secondary (افراد دیگر)',
}
PERIOD_FA = {
    'pre_revolution':  'قبل از انقلاب',
    'early_revolution':'اوایل انقلاب',
    'presidency':      'دوران ریاست جمهوری',
    'hashemi':         'رفسنجانی',
    'khatami':         'خاتمی',
    'ahmadinejad':     'احمدی‌نژاد',
    'rouhani':         'روحانی',
    'raisi_to_death':  'رئیسی تا مرگ',
}

# ─── Permalink: خواندن query_params ─────────────────────────────────────────
_qp = st.query_params
if _qp:
    # فقط اگر ویجت هنوز مقدار session_state ندارد، از URL بخوان
    if 'kw' in _qp and 'kw' not in st.session_state:
        st.session_state['kw'] = _qp['kw']
    if 'nav_y1' not in st.session_state and 'y1' in _qp:
        try: st.session_state['nav_y1'] = int(_qp['y1'])
        except: pass
    if 'nav_y2' not in st.session_state and 'y2' in _qp:
        try: st.session_state['nav_y2'] = int(_qp['y2'])
        except: pass
    if 'nav_period' not in st.session_state and 'period' in _qp:
        st.session_state['nav_period'] = _qp['period']
    if 'nav_genre' not in st.session_state and 'genre' in _qp:
        st.session_state['nav_genre'] = _qp['genre']
    if 'nav_tone' not in st.session_state and 'tone' in _qp:
        st.session_state['nav_tone'] = _qp['tone']
    if 'nav_region' not in st.session_state and 'region' in _qp:
        st.session_state['nav_region'] = _qp['region']

# ─── جستجو ───────────────────────────────────────────────────────────────────
st.markdown("""
<style>
div[data-testid="stTextInput"] > div > div > input {
    font-size: 1.25rem !important;
    padding: 14px 16px !important;
    border: 2px solid #E91E63 !important;
    border-radius: 8px !important;
}
</style>
""", unsafe_allow_html=True)
keyword = st.text_input(
    "کلیدواژه",
    placeholder="کلیدواژه را اینجا بنویسید — مثلاً: آمریکا، مقاومت، هسته‌ای",
    key="kw",
    label_visibility="collapsed",
)
st.caption('برای جستجوی عبارت کامل از گیومه استفاده کنید: «عبارت کامل» | برای حذف کلمه پیشوند منفی بگذارید: ‒کلمه')

with st.expander("تنظیمات جستجو"):
    c1,c2,c3,c4 = st.columns(4)
    logic     = c1.radio("منطق", ["AND","OR"], horizontal=True)
    search_in = c2.radio("جستجو در", ["هر دو","عنوان","متن"], horizontal=True)
    norm_fa   = c3.checkbox("یکسان‌سازی حروف", value=True,
                            help="ي → ی  |  ك → ک  |  ة → ه  |  حذف اعراب")
    use_stem  = c4.checkbox("ریشه‌یابی", value=False,
                            help="پسوندهای فارسی (ها، های، ی، ان، تر، ترین، ...) را حذف و با ریشه جستجو می‌کند")
    c5, c6 = st.columns(2)
    use_wildcard = c5.checkbox("جستجوی wildcard/regex",  value=False,
                                help="* = هر تعداد کاراکتر  |  ? = یک کاراکتر  |  برای regex پیشرفته از .* استفاده کنید\nمثال: استکبار*  یا  انقلاب.اسلامی")
    if use_wildcard:
        c6.caption("🔍 حالت wildcard: * → .* | ? → . | جستجو با regex در Python")
st.divider()

# ─── Sidebar ─────────────────────────────────────────────────────────────────
# ناوبری از نمودارها — قبل از رندر sidebar مقادیر را می‌خوانیم
_nav_period = st.session_state.pop('nav_period', None)
_nav_genre  = st.session_state.pop('nav_genre',  None)
_nav_tone   = st.session_state.pop('nav_tone',   None)
_nav_region = st.session_state.pop('nav_region', None)

with st.sidebar:
    st.header("فیلترها")
    sel_tier    = st.multiselect("نوع سند", opts['tiers'],
                                 format_func=lambda x: TIER_FA.get(x, x))
    sel_src     = st.multiselect("منبع", opts['sources'])
    st.markdown("**تگ‌های AI**")
    sel_voice   = st.multiselect("صدا",     opts['voice'],
                                 format_func=lambda x: VOICE_FA.get(x, x))
    _genre_def  = [_nav_genre]  if _nav_genre  and _nav_genre  in opts['genre']   else []
    sel_genre   = st.multiselect("ژانر",    opts['genre'],   default=_genre_def)
    sel_aud     = st.multiselect("مخاطب",   opts['audience'])
    sel_occ     = st.multiselect("مناسبت",  opts['occasion'])
    sel_topics  = st.multiselect("موضوع",   opts['topics'])
    _region_def = [_nav_region] if _nav_region and _nav_region in opts['regions'] else []
    sel_regions = st.multiselect("منطقه",   opts['regions'], default=_region_def)
    _tone_def   = [_nav_tone]   if _nav_tone   and _nav_tone   in opts['tone']    else []
    sel_tone    = st.multiselect("لحن",     opts['tone'],    default=_tone_def)
    _period_def = [_nav_period] if _nav_period and _nav_period in opts['periods'] else []
    sel_periods = st.multiselect("دوره تاریخی", opts['periods'], default=_period_def,
                                 format_func=lambda x: PERIOD_FA.get(x, x))
    st.markdown("**زمان**")
    nav_y1 = st.session_state.pop('nav_y1', 1368)
    nav_y2 = st.session_state.pop('nav_y2', 1404)
    y1   = st.number_input("از سال", 1360, 1404, int(nav_y1), step=1)
    y2   = st.number_input("تا سال", 1360, 1404, int(nav_y2), step=1)
    wmin = st.number_input("حداقل کلمه", 0, 50000, 0, step=100)
    st.markdown("**نمایش**")
    per_page = st.selectbox("نتایج/صفحه", [10,25,50,100], index=1)
    if st.button("پاک‌کردن"):
        for k in ['nav_y1','nav_y2','nav_kw','nav_period','nav_genre',
                  'nav_tone','nav_region','xl_bytes']:
            st.session_state.pop(k, None)
        st.rerun()

    # ─── ورکست‌های ذخیره‌شده ─────────────────────────────────────────────
    st.divider()
    st.markdown("**📂 ورکست‌های ذخیره‌شده**")
    _ws_list = list_worksets()
    if _ws_list:
        _ws_names = [w['name'] for w in _ws_list]
        _ws_sel = st.selectbox("بارگذاری ورکست", ['—'] + _ws_names, key='ws_sel')
        if _ws_sel != '—' and st.button("بارگذاری", key='ws_load'):
            _ws = get_workset(_ws_sel)
            if _ws and 'filters' in _ws:
                f = _ws['filters']
                for _k, _v in f.items():
                    st.session_state[_k] = _v
                st.rerun()
        if _ws_sel != '—' and st.button("حذف", key='ws_del'):
            delete_workset(_ws_sel)
            st.rerun()
    else:
        st.caption("هنوز ورکستی ذخیره نشده")

    # ذخیره ورکست جدید
    with st.expander("ذخیره ورکست جدید", icon="💾"):
        _ws_name = st.text_input("نام ورکست", key='ws_name_inp',
                                  placeholder="مثال: استکبار_دهه۶۰")
        if st.button("ذخیره", key='ws_save_btn') and _ws_name.strip():
            _ws_filters = {
                'kw': st.session_state.get('kw',''),
                'nav_y1': y1, 'nav_y2': y2,
            }
            save_workset(_ws_name.strip(), _ws_filters)
            st.success(f"ورکست «{_ws_name}» ذخیره شد")
            st.rerun()

# ─── توابع کمکی ──────────────────────────────────────────────────────────────
def normalize_fa(t):
    t = _re.sub(r'[ً-ٰ]','',t)
    return t.replace('ي','ی').replace('ك','ک').replace('ة','ه').replace('‌',' ')

# ریشه‌یاب فارسی (rule-based)
_SUFFIXES = ['هایشان','هایتان','هایمان','هایش','هایت','هایم',
             'های','ترین','شان','تان','مان','گان','ات',
             'ها','تر','ان','ون','ی','م','ت','ش']

def persian_stem(word):
    for s in _SUFFIXES:
        if word.endswith(s) and len(word) - len(s) >= 3:
            return word[:-len(s)]
    return word

def jlist(v):
    try: items = json.loads(v or '[]'); return ', '.join(items) if items else '—'
    except: return v or '—'

def highlight(text, kw):
    if not kw or not text: return _html.escape(text or '')
    escaped = _html.escape(text)
    kw_n = normalize_fa(kw.strip()) if norm_fa else kw.strip()
    for w in [w.lstrip('-') for w in kw_n.split() if w and not w.startswith('-')]:
        if w:
            escaped = _re.compile(_re.escape(_html.escape(w)), _re.IGNORECASE).sub(
                lambda m: f'<mark>{m.group()}</mark>', escaped)
    return escaped

def wildcard_to_like(pattern):
    """تبدیل wildcard به LIKE: * → % و ? → _"""
    result = pattern.replace('%','%%').replace('_',r'\_')
    result = result.replace('*','%').replace('?','_')
    return result

def build_where(keyword, logic, search_in, norm, stem, tiers, sources, voice, genre,
                aud, occ, topics, regions, tone, periods, y1, y2, wmin, use_wildcard=False):
    conds, params = [], []
    if keyword.strip():
        kw = normalize_fa(keyword.strip()) if norm else keyword.strip()
        kw_parts = []
        for w in [x for x in kw.split() if x]:
            if w.startswith('-'):
                ex = w[1:]
                if search_in in ('هر دو','عنوان'): conds.append("d.title NOT LIKE ?"); params.append(f'%{ex}%')
                if search_in in ('هر دو','متن'):   conds.append("d.full_text NOT LIKE ?"); params.append(f'%{ex}%')
                continue
            if use_wildcard:
                # wildcard: * و ? را به SQL LIKE تبدیل می‌کنیم
                sw = wildcard_to_like(w if not stem else persian_stem(w))
                like_pat = sw if '%' in sw or '_' in sw else f'%{sw}%'
            else:
                sw = persian_stem(w) if stem else w
                like_pat = f'%{sw}%'
            col_or = []
            if search_in in ('هر دو','عنوان'): col_or.append("d.title LIKE ?"); params.append(like_pat)
            if search_in in ('هر دو','متن'):   col_or.append("d.full_text LIKE ?"); params.append(like_pat)
            if col_or: kw_parts.append('('+ ' OR '.join(col_or) +')')
        if kw_parts:
            conds.append('(' + (' AND ' if logic=='AND' else ' OR ').join(kw_parts) + ')')
    if tiers:    conds.append(f"d.analytical_tier IN ({','.join('?'*len(tiers))})"); params.extend(tiers)
    if sources:  conds.append(f"d.content_source IN ({','.join('?'*len(sources))})"); params.extend(sources)
    if periods:  conds.append(f"d.period_label IN ({','.join('?'*len(periods))})"); params.extend(periods)
    if voice:    conds.append(f"t.tag_voice_type IN ({','.join('?'*len(voice))})"); params.extend(voice)
    if genre:    conds.append(f"t.tag_form_genre IN ({','.join('?'*len(genre))})"); params.extend(genre)
    if occ:      conds.append(f"t.tag_occasion IN ({','.join('?'*len(occ))})"); params.extend(occ)
    for a  in aud:     conds.append("EXISTS(SELECT 1 FROM json_each(t.tag_audience) WHERE value=?)"); params.append(a)
    for tp in topics:  conds.append("EXISTS(SELECT 1 FROM json_each(t.tag_topics) WHERE value=?)"); params.append(tp)
    for rg in regions: conds.append("EXISTS(SELECT 1 FROM json_each(t.tag_regions) WHERE value=?)"); params.append(rg)
    for tn in tone:    conds.append("EXISTS(SELECT 1 FROM json_each(t.tag_tone) WHERE value=?)"); params.append(tn)
    conds.append("(d.date_persian IS NULL OR d.date_persian BETWEEN ? AND ?)"); params.extend([f"{y1}/01/01", f"{y2}/12/29"])
    if wmin > 0: conds.append("d.word_count >= ?"); params.append(wmin)
    return ("WHERE " + " AND ".join(conds) if conds else ""), params

SORT_OPTIONS = [
    ("date_desc",  "تاریخ: جدید → قدیم"),
    ("date_asc",   "تاریخ: قدیم → جدید"),
    ("words_desc", "کلمات: زیاد → کم"),
    ("words_asc",  "کلمات: کم → زیاد"),
    ("title_asc",  "عنوان: الف → ی"),
    ("title_desc", "عنوان: ی → الف"),
]
ORDER_MAP = {
    "date_desc": "d.date_persian DESC", "date_asc": "d.date_persian ASC",
    "words_desc":"d.word_count DESC",   "words_asc":"d.word_count ASC",
    "title_asc": "d.title ASC",         "title_desc":"d.title DESC",
}

# ─── کوئری ────────────────────────────────────────────────────────────────────
if 'pg' not in st.session_state: st.session_state['pg'] = 0
if keyword != st.session_state.get('_kw',''): st.session_state['pg'] = 0; st.session_state['_kw'] = keyword
# ناوبری از صفحه نمودارها
if 'nav_kw' in st.session_state:
    keyword = st.session_state.pop('nav_kw')
pg = st.session_state['pg']

where, params = build_where(keyword, logic, search_in, norm_fa, use_stem,
    sel_tier, sel_src, sel_voice, sel_genre, sel_aud, sel_occ,
    sel_topics, sel_regions, sel_tone, sel_periods, y1, y2, wmin,
    use_wildcard=use_wildcard)
join = "LEFT JOIN doc_tags t ON d.doc_id=t.doc_id"

try:
    total = conn.execute(f"SELECT COUNT(*) FROM documents d {join} {where}", params).fetchone()[0]
except Exception as e:
    st.error(f"خطای جستجو: {e}"); st.stop()

# ─── نوار ابزار ──────────────────────────────────────────────────────────────
tb1, tb2, tb3, tb4 = st.columns([3, 2, 1.2, 1.2])
tb1.markdown(f'<div class="rtl" style="padding-top:8px"><b>{total:,} سند یافت شد</b> — صفحه {pg+1} از {max(1,(total+per_page-1)//per_page)}</div>',
             unsafe_allow_html=True)
sort = tb2.selectbox("مرتب‌سازی", SORT_OPTIONS, format_func=lambda x: x[1], label_visibility="collapsed")
show_charts = tb3.checkbox("نمودار نتایج", key="show_charts")
make_xl = tb4.button("Excel همه نتایج", use_container_width=True)

# ─── Permalink ───────────────────────────────────────────────────────────────
with st.expander("پیوند اشتراک‌گذاری", icon="🔗", expanded=False):
    import urllib.parse as _up
    _base = "http://localhost:8502/%D8%AC%D8%B3%D8%AA%D8%AC%D9%88"
    _pl_params = {}
    if keyword.strip():     _pl_params['kw']     = keyword.strip()
    if y1 != 1356:          _pl_params['y1']     = str(y1)
    if y2 != 1404:          _pl_params['y2']     = str(y2)
    if sel_tier:            _pl_params['tier']   = ','.join(sel_tier)
    if sel_src:             _pl_params['src']    = ','.join(sel_src)
    if sel_periods:         _pl_params['period'] = ','.join(sel_periods)
    if sel_genre:           _pl_params['genre']  = ','.join(sel_genre)
    if sel_tone:            _pl_params['tone']   = ','.join(sel_tone)
    if sel_regions:         _pl_params['region'] = ','.join(sel_regions)
    _qs = _up.urlencode(_pl_params, quote_via=_up.quote)
    _permalink = f"{_base}?{_qs}" if _qs else _base
    st.code(_permalink, language=None)
    st.caption("این لینک را کپی کنید — کلیدواژه، فیلترها و بازه سال را ذخیره می‌کند. برای تنظیم آدرس سرور، پارامتر base URL را در کد تغییر دهید.")

# ─── Excel ───────────────────────────────────────────────────────────────────
@st.cache_data(ttl=180, show_spinner=False)
def _build_excel(where_str, params_t, sort_key, kw, y1v, y2v):
    import sqlite3 as _sq
    from openpyxl import Workbook
    from openpyxl.styles import Font as XFont, PatternFill, Alignment as XAlign
    _c = _sq.connect(DB_PATH); _c.row_factory = _sq.Row
    _join = "LEFT JOIN doc_tags t ON d.doc_id=t.doc_id"
    _rows = [dict(r) for r in _c.execute(
        f"SELECT d.doc_id,d.date_persian,d.analytical_tier,d.content_source,"
        f"d.period_label,d.title,d.word_count,d.url,d.full_text,"
        f"t.tag_voice_type,t.tag_form_genre,t.tag_occasion,"
        f"t.tag_topics,t.tag_regions,t.tag_audience,t.tag_tone "
        f"FROM documents d {_join} {where_str} "
        f"ORDER BY {ORDER_MAP.get(sort_key,'d.date_persian DESC')} LIMIT 10000",
        list(params_t)).fetchall()]
    _c.close()
    wb = Workbook()
    ws0 = wb.active; ws0.title='اطلاعات'; ws0.sheet_view.rightToLeft=True
    for r in [['کلیدواژه', kw or '—'],['تعداد نتایج', len(_rows)],
               ['بازه سال', f"{y1v}–{y2v}"],['تاریخ', datetime.now().strftime('%Y-%m-%d %H:%M')]]:
        ws0.append(r)
    ws1 = wb.create_sheet('متادیتا'); ws1.sheet_view.rightToLeft=True
    mc = ['doc_id','date_persian','analytical_tier','content_source','period_label',
          'title','word_count','url','tag_voice_type','tag_form_genre',
          'tag_occasion','tag_topics','tag_regions','tag_audience','tag_tone']
    ws1.append(mc)
    for row in _rows: ws1.append([row.get(c,'') for c in mc])
    ws2 = wb.create_sheet('متن کامل'); ws2.sheet_view.rightToLeft=True
    ws2.append(['doc_id','date_persian','title','full_text'])
    for row in _rows: ws2.append([row.get('doc_id',''),row.get('date_persian',''),row.get('title',''),row.get('full_text','')])
    ws3 = wb.create_sheet('تگ ۷ بعدی'); ws3.sheet_view.rightToLeft=True
    tc = ['doc_id','tag_voice_type','tag_form_genre','tag_audience','tag_occasion','tag_topics','tag_regions','tag_tone']
    ws3.append(tc)
    for row in _rows: ws3.append([row.get(c,'') for c in tc])
    hf = XFont(bold=True, name='Vazirmatn'); hfill = PatternFill('solid', fgColor='E8E8E8')
    for ws in [ws0,ws1,ws2,ws3]:
        for cell in ws[1]: cell.font=hf; cell.fill=hfill
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.font=XFont(name='Vazirmatn',size=9)
                cell.alignment=XAlign(wrap_text=True,vertical='top')
    buf = io.BytesIO(); wb.save(buf)
    return buf.getvalue(), len(_rows)

if make_xl:
    with st.spinner("در حال ساخت Excel…"):
        try:
            xl_b, xl_n = _build_excel(where, tuple(params), sort[0], keyword, y1, y2)
            st.session_state['xl_bytes'] = xl_b
            st.session_state['xl_name']  = f"khamenei_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
            st.session_state['xl_count'] = xl_n
        except Exception as e:
            st.error(f"خطای Excel: {e}")

if st.session_state.get('xl_bytes'):
    xl_col1, xl_col2 = st.columns([3,1])
    xl_col1.success(f"{st.session_state.get('xl_count',0):,} سند آماده است")
    xl_col2.download_button("دانلود Excel",
        st.session_state['xl_bytes'],
        file_name=st.session_state.get('xl_name','khamenei.xlsx'),
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        key="dl_xl_final", use_container_width=True)

# ─── ساخت PDF ─────────────────────────────────────────────────────────────────
def make_pdf(row_data):
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        from fpdf import FPDF

        def r(text):
            if not text: return ''
            return get_display(arabic_reshaper.reshape(str(text)))

        pdf = FPDF()
        pdf.add_page()
        pdf.add_font('V', fname=FONT_PATH)
        pdf.set_right_margin(15); pdf.set_left_margin(15)

        # عنوان
        pdf.set_font('V', size=16)
        pdf.multi_cell(0, 10, r(row_data.get('title','(بدون عنوان)')), align='R')
        pdf.ln(2)

        # متادیتا
        pdf.set_font('V', size=10)
        pdf.set_text_color(100,100,100)
        meta = f"{row_data.get('date_persian','')} | {row_data.get('content_source','')} | {row_data.get('analytical_tier','')}"
        pdf.multi_cell(0, 7, r(meta), align='R')
        if row_data.get('url'):
            pdf.multi_cell(0, 7, r(f"لینک: {row_data.get('url','')}"), align='R')
        pdf.ln(3)
        pdf.set_text_color(0,0,0)

        # خط جدا
        pdf.set_draw_color(200,200,200)
        pdf.line(15, pdf.get_y(), 195, pdf.get_y())
        pdf.ln(4)

        # تگ‌ها
        pdf.set_font('V', size=10)
        pdf.set_fill_color(245,245,245)
        tags = [
            f"voice: {row_data.get('tag_voice_type','—')}  |  genre: {row_data.get('tag_form_genre','—')}",
            f"audience: {jlist(row_data.get('tag_audience'))}",
            f"topics: {jlist(row_data.get('tag_topics'))}",
            f"regions: {jlist(row_data.get('tag_regions'))}",
            f"tone: {jlist(row_data.get('tag_tone'))}",
        ]
        for line in tags:
            pdf.multi_cell(0, 7, r(line), align='R', fill=True)
        pdf.ln(3)
        pdf.line(15, pdf.get_y(), 195, pdf.get_y())
        pdf.ln(4)

        # متن کامل
        pdf.set_font('V', size=11)
        ft = row_data.get('full_text','') or ''
        # به پاراگراف‌ها تقسیم
        for para in ft.split('\n'):
            if para.strip():
                pdf.multi_cell(0, 8, r(para.strip()), align='R')
                pdf.ln(1)

        pdf.set_font('V', size=8)
        pdf.set_text_color(150,150,150)
        pdf.cell(0, 6, r(f"پیکره خامنه‌ای — {datetime.now().strftime('%Y-%m-%d')}"), align='C')

        buf = io.BytesIO()
        pdf.output(buf)
        return buf.getvalue()
    except Exception as e:
        return None

# ─── نمودار نتایج ─────────────────────────────────────────────────────────────
if show_charts:
    FONT_C = dict(family="Vazirmatn, sans-serif", size=12)
    BG     = dict(plot_bgcolor='#fff', paper_bgcolor='#fff')
    ts_c   = datetime.now().strftime('%H%M%S')

    with st.expander("نمودارهای نتایج جستجو", expanded=True):

        def dl_chart(fig, name, df_csv=None):
            ca, cb = st.columns(2)
            if df_csv is not None:
                ca.download_button("دانلود CSV", df_csv.to_csv(index=False).encode('utf-8-sig'),
                                   file_name=f"{name}_{ts_c}.csv", key=f"csv_{name}_{ts_c}")
            cb.download_button("دانلود JSON", fig.to_json().encode(),
                               file_name=f"{name}_{ts_c}.json", key=f"jsn_{name}_{ts_c}")
            st.caption("PNG: روی نمودار hover کنید → دکمه دوربین")

        # ۱. سالانه
        yr_rows = conn.execute(f"""
            SELECT substr(d.date_persian,1,4) y, COUNT(*) n FROM documents d {join} {where}
            AND d.date_persian GLOB '1[34][0-9][0-9]/*' GROUP BY y
            HAVING CAST(y AS INTEGER) BETWEEN 1356 AND 1405 ORDER BY y
        """, params).fetchall()
        if yr_rows:
            st.subheader("توزیع سالانه")
            df_yr = pd.DataFrame(yr_rows, columns=['سال','تعداد'])
            fig_yr = px.bar(df_yr, x='سال', y='تعداد', color_discrete_sequence=['#E91E63'],
                            title=f'توزیع سالانه — {keyword or "همه"}')
            fig_yr.update_layout(font=FONT_C, **BG, height=320, margin=dict(l=40,r=20,t=50,b=60),
                                  xaxis=dict(range=['1356', '1406'], tickangle=-45))
            st.plotly_chart(fig_yr, use_container_width=True)
            dl_chart(fig_yr, 'yearly', df_yr)
        st.divider()

        # ۲. ژانر (پای tier حذف شد)
        gn_rows = conn.execute(f"""
            SELECT t.tag_form_genre g, COUNT(*) n FROM documents d {join} {where}
            AND t.tag_form_genre IS NOT NULL GROUP BY g ORDER BY n DESC LIMIT 12
        """, params).fetchall()
        if gn_rows:
            df_g = pd.DataFrame(gn_rows, columns=['ژانر','تعداد'])
            df_g = df_g.sort_values('تعداد')
            fig_g = px.bar(df_g, x='تعداد', y='ژانر', orientation='h', title='توزیع ژانر',
                           color='تعداد', color_continuous_scale='Reds', text='تعداد')
            fig_g.update_traces(textposition='outside')
            fig_g.update_layout(font=FONT_C, **BG, height=380, margin=dict(l=200,r=80,t=50,b=20))
            fig_g.update_yaxes(automargin=True)
            st.plotly_chart(fig_g, use_container_width=True)
            with st.expander("دانلود داده ژانر"):
                dl_chart(fig_g, 'genre', df_g)
        st.divider()

        # ۳. لحن + دوره
        col_c, col_d = st.columns(2)
        tone_rows = conn.execute(f"""
            SELECT jt.value, COUNT(*) n FROM documents d {join}
            JOIN json_each(t.tag_tone) jt {where}
            AND jt.value != '' GROUP BY jt.value ORDER BY n DESC LIMIT 12
        """, params).fetchall()
        if tone_rows:
            df_tn = pd.DataFrame(tone_rows, columns=['لحن','تعداد'])
            df_tn = df_tn.sort_values('تعداد')
            fig_tn = px.bar(df_tn, x='تعداد', y='لحن', orientation='h', title='لحن',
                            color='تعداد', color_continuous_scale='Blues', text='تعداد')
            fig_tn.update_traces(textposition='outside')
            fig_tn.update_layout(font=FONT_C, **BG, height=350, margin=dict(l=180,r=60,t=50,b=20))
            col_c.plotly_chart(fig_tn, use_container_width=True)

        pr_rows = conn.execute(f"""
            SELECT d.period_label, COUNT(*) n FROM documents d {join} {where}
            AND d.period_label IS NOT NULL GROUP BY d.period_label ORDER BY MIN(d.date_persian)
        """, params).fetchall()
        if pr_rows:
            df_pr = pd.DataFrame(pr_rows, columns=['دوره','تعداد'])
            df_pr['دوره'] = df_pr['دوره'].map(PERIOD_FA).fillna(df_pr['دوره'])
            fig_pr = px.bar(df_pr, x='تعداد', y='دوره', orientation='h', title='دوره',
                            color='تعداد', color_continuous_scale='Greens', text='تعداد')
            fig_pr.update_traces(textposition='outside')
            fig_pr.update_layout(font=FONT_C, **BG, height=320,
                                 margin=dict(l=180,r=60,t=50,b=20), yaxis=dict(autorange='reversed'))
            col_d.plotly_chart(fig_pr, use_container_width=True)

        with st.expander("دانلود داده لحن / دوره"):
            if tone_rows: dl_chart(fig_tn, 'tone', df_tn)
            if pr_rows:   dl_chart(fig_pr, 'period', df_pr)
        st.divider()

        # ۴. مقایسه کلیدواژه
        st.subheader("مقایسه کلیدواژه‌ها")
        st.caption("تا ۶ کلیدواژه با کاما جدا کنید")
        cmp_raw = st.text_input("کلیدواژه‌ها", value=keyword or "", key="cmp_kw")
        if cmp_raw.strip():
            cmp_kws = [k.strip() for k in cmp_raw.split(',') if k.strip()][:6]
            fig_cmp = go.Figure()
            colors = ['#E91E63','#2196F3','#4CAF50','#FF9800','#9C27B0','#00BCD4']
            cmp_data = {}
            for i, kw_c in enumerate(cmp_kws):
                trend = conn.execute("""
                    SELECT substr(date_persian,1,4) y, COUNT(*) n FROM documents
                    WHERE date_persian GLOB '1[34][0-9][0-9]/*' AND full_text LIKE ?
                    GROUP BY y ORDER BY y
                """, (f'%{kw_c}%',)).fetchall()
                if trend:
                    yrs=[r[0] for r in trend]; cnts=[r[1] for r in trend]
                    cmp_data[kw_c] = dict(zip(yrs, cnts))
                    fig_cmp.add_trace(go.Scatter(x=yrs, y=cnts, mode='lines+markers',
                        name=kw_c, line=dict(color=colors[i%6], width=2), marker=dict(size=5)))
            if fig_cmp.data:
                fig_cmp.update_layout(font=FONT_C, **BG, height=380,
                                      margin=dict(l=60,r=20,t=50,b=60),
                                      xaxis_title='سال شمسی', yaxis_title='تعداد سند',
                                      hovermode='x unified',
                                      title=f"مقایسه: {' / '.join(cmp_kws)}")
                st.plotly_chart(fig_cmp, use_container_width=True)
                if cmp_data:
                    all_years = sorted(set(y for d in cmp_data.values() for y in d))
                    df_cmp = pd.DataFrame({'سال': all_years})
                    for kw_c, yd in cmp_data.items():
                        df_cmp[kw_c] = df_cmp['سال'].map(yd).fillna(0).astype(int)
                    dl_chart(fig_cmp, 'comparison', df_cmp)

# ─── واکشی صفحه ──────────────────────────────────────────────────────────────
order_sql = ORDER_MAP.get(sort[0], "d.date_persian DESC")
try:
    rows = [dict(r) for r in conn.execute(f"""
        SELECT d.doc_id, d.date_persian, d.title, d.word_count,
               d.content_source, d.analytical_tier, d.period_label, d.url,
               substr(d.full_text,1,400) AS snippet,
               t.tag_voice_type, t.tag_form_genre, t.tag_occasion,
               t.tag_topics, t.tag_regions, t.tag_audience, t.tag_tone
        FROM documents d {join} {where}
        ORDER BY {order_sql} LIMIT ? OFFSET ?
    """, params + [per_page, pg*per_page]).fetchall()]
except Exception as e:
    st.error(f"خطای بازیابی: {e}"); st.stop()

if not rows:
    st.info("نتیجه‌ای یافت نشد."); st.stop()

st.divider()

# ─── نمایش نتایج ─────────────────────────────────────────────────────────────
for row in rows:
    tier  = row.get('analytical_tier','') or ''
    voice = row.get('tag_voice_type','') or ''
    tc = 'chip chip-t1' if tier=='tier1' else 'chip'
    vc = 'chip chip-pri' if voice=='primary' else 'chip'
    snippet_hl = highlight(row.get('snippet','')[:280], keyword)
    title_hl   = highlight(row.get('title') or '(بدون عنوان)', keyword)

    st.markdown(f"""
<div class="card">
  <div class="card-title">{title_hl}</div>
  <div class="card-meta">{row.get('date_persian','')} &nbsp;|&nbsp; {row.get('content_source','')} &nbsp;|&nbsp; {PERIOD_FA.get(row.get('period_label',''), row.get('period_label',''))}</div>
  <div class="card-snippet">{snippet_hl}…</div>
  <div>
    <span class="{tc}">{TIER_FA.get(tier, tier) or '—'}</span>
    <span class="{vc}">{VOICE_FA.get(voice, voice) or '—'}</span>
    <span class="chip">{row.get('tag_form_genre','') or '—'}</span>
    <span class="chip">{row.get('word_count',0):,} کلمه</span>
  </div>
</div>""", unsafe_allow_html=True)

    if row.get('url'):
        st.markdown(f"[سایت رسمی]({row['url']})")

    with st.expander("‏جزئیات و متن کامل"):
        doc_id = row['doc_id']
        full_r = conn.execute("""
            SELECT d.doc_id, d.date_persian, d.title, d.word_count,
                   d.content_source, d.analytical_tier, d.period_label, d.url,
                   d.full_text,
                   t.tag_voice_type, t.tag_form_genre, t.tag_occasion,
                   t.tag_topics, t.tag_regions, t.tag_audience, t.tag_tone
            FROM documents d LEFT JOIN doc_tags t ON d.doc_id=t.doc_id
            WHERE d.doc_id=?
        """, (doc_id,)).fetchone()

        if full_r:
            r = dict(full_r)
            st.markdown(f"""
**تگ‌های ۷ بعدی:**
- صدا: `{VOICE_FA.get(r.get('tag_voice_type',''), r.get('tag_voice_type','—'))}` | ژانر: `{r.get('tag_form_genre','—')}` | مناسبت: `{r.get('tag_occasion') or '—'}`
- مخاطب: `{jlist(r.get('tag_audience'))}`
- موضوع: `{jlist(r.get('tag_topics'))}`
- منطقه: `{jlist(r.get('tag_regions'))}`
- لحن: `{jlist(r.get('tag_tone'))}`
""")
            # ─── Citation Generator ───────────────────────────────────────────
            with st.expander("استناددهی (Citation Generator)", icon="📌"):
                _title   = r.get('title','') or '(بدون عنوان)'
                _date    = r.get('date_persian','') or ''
                _url     = r.get('url','') or ''
                _src     = r.get('content_source','') or ''
                _tier    = r.get('analytical_tier','') or ''
                _doc_id  = r.get('doc_id','')
                # تبدیل تاریخ شمسی به میلادی تقریبی
                _year_g  = ''
                try:
                    _ysh = int(_date[:4])
                    _year_g = str(_ysh + 621)
                except: pass

                _apa = (
                    f"Khamenei, A. (‎{_year_g or _date[:4]}). "
                    f"*{_title}* [{_src}]. "
                    f"Khamenei Corpus, Volant Media. "
                    f"{_url}"
                )
                _chicago = (
                    f'Khamenei, Ali. "{_title}." '
                    f'{_src.capitalize()}, {_year_g or _date[:4]}. '
                    f'Khamenei Corpus (Volant Media). '
                    f'{_url}.'
                )
                _fa = (
                    f"خامنه‌ای، علی. «{_title}». "
                    f"{_date[:7] if _date else ''}. "
                    f"پیکره گفتاری خامنه‌ای، لانت مدیا. "
                    f"بازیابی از: {_url}"
                )
                _bibtex = (
                    f"@misc{{{_doc_id},\n"
                    f"  author = {{Khamenei, Ali}},\n"
                    f"  title  = {{{_title}}},\n"
                    f"  year   = {{{_year_g or _date[:4]}}},\n"
                    f"  howpublished = {{\\url{{{_url}}}}},\n"
                    f"  note   = {{Khamenei Corpus, Volant Media; {_tier}}}\n"
                    f"}}"
                )

                ct1, ct2, ct3, ct4 = st.tabs(["APA", "Chicago", "فارسی", "BibTeX"])
                with ct1: st.code(_apa,     language=None)
                with ct2: st.code(_chicago, language=None)
                with ct3: st.code(_fa,      language=None)
                with ct4: st.code(_bibtex,  language=None)

            st.divider()
            ft = r.get('full_text','') or ''
            ft_hl = highlight(ft, keyword)
            st.markdown(
                f'<div style="height:380px;overflow-y:scroll;direction:rtl;line-height:1.9;'
                f'font-size:0.92rem;padding:10px;background:#fafafa;border-radius:6px">{ft_hl}</div>',
                unsafe_allow_html=True)

            st.markdown("**دانلود سند:**")
            dl1, dl2, dl3 = st.columns(3)
            dl1.download_button("متن (.txt)", ft.encode('utf-8'),
                                file_name=f"{doc_id}.txt", mime="text/plain",
                                key=f"txt_{doc_id}", use_container_width=True)

            try:
                from docx import Document as DocxDoc
                from docx.shared import Pt
                d = DocxDoc()
                hp = d.add_heading(r.get('title','') or '', 0); hp.alignment = 2
                mp = d.add_paragraph()
                mp.add_run(f"تاریخ: {r.get('date_persian','')} | Tier: {r.get('analytical_tier','')}").font.size = Pt(10)
                mp.alignment = 2
                d.add_heading('تگ‌ها', 2)
                tp2 = d.add_paragraph()
                for label, key2 in [('voice','tag_voice_type'),('genre','tag_form_genre'),
                                    ('topics','tag_topics'),('regions','tag_regions'),('tone','tag_tone')]:
                    run = tp2.add_run(f"{label}: "); run.bold = True; run.font.size = Pt(10)
                    tp2.add_run(str(r.get(key2,'')) + '\n').font.size = Pt(10)
                tp2.alignment = 2
                d.add_heading('متن', 2)
                bp = d.add_paragraph(ft); bp.alignment = 2
                wbuf = io.BytesIO(); d.save(wbuf)
                dl2.download_button("Word (.docx)", wbuf.getvalue(),
                                    file_name=f"{doc_id}.docx",
                                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                                    key=f"docx_{doc_id}", use_container_width=True)
            except Exception:
                dl2.caption("Word: python-docx نیاز دارد")

            pdf_bytes = make_pdf(r)
            if pdf_bytes:
                dl3.download_button("PDF", pdf_bytes,
                                    file_name=f"{doc_id}.pdf", mime="application/pdf",
                                    key=f"pdf_{doc_id}", use_container_width=True)
            else:
                dl3.caption("PDF در دسترس نیست")

# ─── صفحه‌بندی ────────────────────────────────────────────────────────────────
total_pages = max(1, (total + per_page - 1) // per_page)
if total_pages > 1:
    st.divider()
    p1, p2, p3 = st.columns([1,3,1])
    if pg > 0 and p1.button("قبلی"):
        st.session_state['pg'] = pg-1; st.rerun()
    p2.markdown(f'<div style="text-align:center">صفحه {pg+1} از {total_pages}</div>', unsafe_allow_html=True)
    if pg < total_pages-1 and p3.button("بعدی"):
        st.session_state['pg'] = pg+1; st.rerun()

render_footer()
